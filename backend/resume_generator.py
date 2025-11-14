# resume_generator.py

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors

def generate_resume(resume_data, template="modern", output_folder="static/generated"):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"resume_{timestamp}"
    
    # --- CRITICAL FIX: Data Normalization ---
    # This ensures we always have valid data dictionaries, even if empty
    data = {
        "personal_info": resume_data.get("personal_info", {}),
        "summary": resume_data.get("summary", ""),
        "skills": resume_data.get("skills", []),
        "experience": resume_data.get("experience", []),
        "education": resume_data.get("education", []),
        "projects": resume_data.get("projects", []),
        "certifications": resume_data.get("certifications", [])
    }
    
    pdf_path = generate_pdf(data, template, output_folder, base_filename)
    docx_path = generate_docx(data, template, output_folder, base_filename)
    
    return {
        "pdf": os.path.basename(pdf_path),
        "docx": os.path.basename(docx_path)
    }

def generate_pdf(data, template, output_folder, base_filename):
    filepath = os.path.join(output_folder, f"{base_filename}.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Define Styles
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, alignment=1, spaceAfter=10)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, spaceBefore=12, spaceAfter=6, 
                              borderWidth=1, borderColor=colors.HexColor('#2c3e50'), borderPadding=5, borderRadius=2)
    normal_style = styles['Normal']
    
    # 1. Personal Info
    p_info = data["personal_info"]
    name = p_info.get("name", "Your Name")
    story.append(Paragraph(name, title_style))
    
    contact_line = [
        p_info.get("email"), 
        p_info.get("phone"), 
        p_info.get("location"),
        p_info.get("linkedin")
    ]
    # Filter out None or empty strings
    contact_text = " | ".join([c for c in contact_line if c])
    story.append(Paragraph(contact_text, ParagraphStyle('Contact', parent=normal_style, alignment=1)))
    story.append(Spacer(1, 0.2*inch))

    # 2. Summary
    if data["summary"]:
        story.append(Paragraph("PROFESSIONAL SUMMARY", h2_style))
        story.append(Paragraph(data["summary"], normal_style))

    # 3. Skills
    if data["skills"]:
        story.append(Paragraph("SKILLS", h2_style))
        # Handle both list of strings and complex objects
        skill_text = ", ".join(data["skills"]) if isinstance(data["skills"][0], str) else ", ".join([s.get("name", "") for s in data["skills"]])
        story.append(Paragraph(skill_text, normal_style))

    # 4. Experience
    if data["experience"]:
        story.append(Paragraph("EXPERIENCE", h2_style))
        for exp in data["experience"]:
            # Safety checks for missing keys
            title = exp.get("title", "Job Title")
            company = exp.get("company", "Company")
            duration = exp.get("duration", "")
            
            header = f"<b>{title}</b> at {company}"
            if duration:
                header += f" ({duration})"
                
            story.append(Paragraph(header, normal_style))
            
            # Handle description list
            descriptions = exp.get("description", [])
            if isinstance(descriptions, str): descriptions = [descriptions]
            
            for desc in descriptions:
                story.append(Paragraph(f"• {desc}", ParagraphStyle('Bullet', parent=normal_style, leftIndent=15)))
            story.append(Spacer(1, 0.1*inch))

    # 5. Education
    if data["education"]:
        story.append(Paragraph("EDUCATION", h2_style))
        for edu in data["education"]:
            degree = edu.get("degree", "Degree")
            school = edu.get("institution", "University")
            year = edu.get("year", "")
            text = f"<b>{degree}</b> - {school} {f'({year})' if year else ''}"
            story.append(Paragraph(text, normal_style))

    doc.build(story)
    return filepath

def generate_docx(data, template, output_folder, base_filename):
    filepath = os.path.join(output_folder, f"{base_filename}.docx")
    doc = Document()
    
    # Personal Info
    p_info = data["personal_info"]
    name = p_info.get("name", "Your Name")
    heading = doc.add_paragraph(name)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.bold = True
    run.font.size = Pt(24)
    
    contact_line = [p_info.get("email"), p_info.get("phone"), p_info.get("linkedin")]
    contact_text = " | ".join([c for c in contact_line if c])
    contact_para = doc.add_paragraph(contact_text)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Helper to add sections
    def add_section(title, content_list, is_list=False):
        if not content_list: return
        h = doc.add_heading(title, level=1)
        
        if title == "SKILLS":
            doc.add_paragraph(", ".join(content_list))
        elif title == "PROFESSIONAL SUMMARY":
            doc.add_paragraph(content_list)
        else:
            for item in content_list:
                if title == "EXPERIENCE":
                    p = doc.add_paragraph()
                    p.add_run(f"{item.get('title', '')} at {item.get('company', '')}").bold = True
                    if item.get('duration'):
                        p.add_run(f" ({item['duration']})").italic = True
                    
                    descs = item.get('description', [])
                    if isinstance(descs, str): descs = [descs]
                    for d in descs:
                        doc.add_paragraph(d, style='List Bullet')
                
                elif title == "EDUCATION":
                    p = doc.add_paragraph()
                    p.add_run(f"{item.get('degree', '')}").bold = True
                    p.add_run(f" - {item.get('institution', '')}")

    add_section("PROFESSIONAL SUMMARY", data["summary"])
    add_section("SKILLS", data["skills"])
    add_section("EXPERIENCE", data["experience"])
    add_section("EDUCATION", data["education"])

    doc.save(filepath)
    return filepath